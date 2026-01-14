from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def main():
    document = Document()

    # Title
    title = document.add_heading('CraveCart - Project Report Documentation', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 1. Introduction
    document.add_heading('1. Introduction', level=1)
    
    document.add_heading('a. Project Profile', level=2)
    p = document.add_paragraph()
    p.add_run('Project Name: ').bold = True
    p.add_run('CraveCart - Food Delivery System\n')
    p.add_run('Type: ').bold = True
    p.add_run('Web Application\n')
    p.add_run('Developers: ').bold = True
    p.add_run('Riya Jogi & Diya Vaghasiya\n')
    p.add_run('Domain: ').bold = True
    p.add_run('E-Commerce / Food Industry')

    document.add_heading('b. Project Abstract', level=2)
    document.add_paragraph(
        "CraveCart is an online food delivery platform designed to connect users with a variety of food items "
        "available for purchase. The application aims to simplify the process of browsing food menus, managing "
        "shopping carts, and placing orders securely. It provides a user-friendly interface for customers to "
        "track their orders and an admin panel for administrators to manage categories, products, and order statuses, "
        "efficient and digitalizing the traditional food ordering process."
    )

    # 2. Analysis
    document.add_heading('2. Analysis', level=1)

    document.add_heading('a. Requirement Analysis', level=2)
    p = document.add_paragraph("The core requirements identified for the system are:")
    items = [
        "User Registration/Login: A secure system for users to access their accounts.",
        "Product Catalog: A categorized display of food items with images and prices.",
        "Shopping Cart: Persistent cart management allowing users to add/remove items and adjust quantities.",
        "Order Processing: A streamlined checkout flow capturing shipping details and payment methods.",
        "Order Tracking: Visibility into the status of placed orders.",
        "Administration: Backend management for inventory and orders."
    ]
    for item in items:
        document.add_paragraph(item, style='List Bullet')

    document.add_heading('b. Current System', level=2)
    document.add_paragraph(
        "The current system typically involves manual phone ordering or visiting restaurants physically. "
        "This process is time-consuming, lacks real-time menu updates, and offers no immediate tracking mechanism. "
        "It is prone to human error in order taking and lacks a digital record for users to review past orders."
    )

    document.add_heading('c. Proposed System', level=2)
    document.add_paragraph(
        "The proposed CraveCart system automates the entire flow. Users can view the latest menu instantly, "
        "place orders without phone calls, and receive visual confirmation of their orders. It reduces manual "
        "errors, saves time, and provides a centralized platform for managing food delivery operations."
    )

    document.add_heading('d. Feasibility Study', level=2)
    document.add_paragraph(
        "Technical Feasibility: Python and Django are robust technologies well-suited for building scalable web applications.\n"
        "Operational Feasibility: The user interface is intuitive, requiring minimal training for users.\n"
        "Economic Feasibility: Built with open-source technologies (Django, SQLite), the development cost is low."
    )

    # 3. Back-End
    document.add_heading('3. About Back-End', level=1)
    document.add_paragraph(
        "The backend is powered by Django, a high-level Python web framework.\n"
        "Language: Python\n"
        "Framework: Django 5.0\n"
        "Database: SQLite\n"
        "Key Features: Django's ORM, User Authentication system, and Admin Interface."
    )

    # 4. Front-End
    document.add_heading('4. About Front-End', level=1)
    document.add_paragraph(
        "Languages: HTML5, CSS3, JavaScript\n"
        "Framework: Bootstrap / Custom CSS\n"
        "Templating: Django Template Language (DTL)"
    )

    # 5. Tools
    document.add_heading('5. About Other Tools Used', level=1)
    document.add_paragraph(
        "VS Code: Primary code editor.\n"
        "Git: Version control system.\n"
        "Pip: Python package manager.\n"
        "Web Browser: Chrome/Firefox for testing."
    )

    # 6. SRS
    document.add_heading('6. Software Requirement & Specification', level=1)
    
    document.add_heading('a. Introduction', level=2)
    document.add_paragraph("This section details the functional and non-functional requirements.")

    document.add_heading('b. Functional Requirements', level=2)
    document.add_paragraph("1. User Authentication: Sign up, Login, Logout.")
    document.add_paragraph("2. Product Management: View categories, items, prices.")
    document.add_paragraph("3. Cart Methodology: Add/Remove items, update quantity.")
    document.add_paragraph("4. Order System: Shipping details, Payment methods, Order IDs.")
    document.add_paragraph("5. Administrative Functions: Manage Inventory, View Orders.")

    document.add_heading('c. Hardware/Software Requirement', level=2)
    document.add_paragraph("Software: OS (Windows/Mac/Linux), Python 3.8+, Django 5.0+, Chrome.")
    document.add_paragraph("Hardware: Intel Core i3+, 4GB RAM, 500MB Storage.")

    # 7. Data Dictionary
    document.add_heading('7. Data Dictionary', level=1)

    def add_table(doc, title, headers, rows):
        doc.add_heading(title, level=3)
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        for i, h in enumerate(headers):
            hdr_cells[i].text = h
            hdr_cells[i].paragraphs[0].runs[0].bold = True
        
        for row_data in rows:
            row_cells = table.add_row().cells
            for i, item in enumerate(row_data):
                row_cells[i].text = str(item)

    # User Table
    add_table(document, 'Table: User (Django Built-in)', 
              ['Field', 'Type', 'Description'],
              [
                  ['id', 'Integer', 'Primary Key'],
                  ['username', 'Varchar', 'Unique username'],
                  ['password', 'Varchar', 'Hashed password'],
                  ['email', 'Varchar', 'User email address']
              ])

    # Category Table
    add_table(document, 'Table: Category', 
              ['Field', 'Type', 'Description'],
              [
                  ['id', 'Integer', 'Primary Key'],
                  ['name', 'Varchar(100)', 'Name of the food category (Unique)']
              ])

    # FoodItem Table
    add_table(document, 'Table: FoodItem', 
              ['Field', 'Type', 'Description'],
              [
                  ['id', 'Integer', 'Primary Key'],
                  ['name', 'Varchar(100)', 'Name of the food item'],
                  ['price', 'Decimal(10,2)', 'Cost of the item'],
                  ['category_id', 'ForeignKey', 'Link to Category table'],
                  ['is_available', 'Boolean', 'Availability status']
              ])

    # Order Table
    add_table(document, 'Table: Order', 
              ['Field', 'Type', 'Description'],
              [
                  ['id', 'Integer', 'Primary Key'],
                  ['user_id', 'ForeignKey', 'Link to User table'],
                  ['total_price', 'Decimal(10,2)', 'Total cost of the order'],
                  ['status', 'Varchar', 'Order status'],
                  ['payment_method', 'Varchar', 'Selected payment method']
              ])

    # 8. E-R Diagram
    document.add_heading('8. E-R Diagram', level=1)
    document.add_paragraph("[Mermaid Diagram provided in Source Report - Visual representation requires external tool]")
    document.add_paragraph("Relationships:\n- User places Order\n- User has Cart\n- Category contains FoodItem\n- Order contains OrderItem\n- FoodItem is in OrderItem/Cart")

    # 9. DFD
    document.add_heading('9. Data Flow Diagram', level=1)
    document.add_paragraph("Level 0 (Context): User -> System -> Admin")
    document.add_paragraph("Level 1 (Flow): Auth -> Browse -> Cart -> Checkout -> Order Processing -> Database")

    # 10. Screen Layouts
    document.add_heading('10. Screen Layouts', level=1)
    layouts = [
        "Home Page: Displays Navbar, Hero Section, and Categories/Food Items Grid.",
        "Login/Signup: Simple form centered on screen.",
        "Cart: Table view list of added items, quantity controls, line totals.",
        "Checkout: Form for shipping details and Payment selection.",
        "Order History: List of past orders showing Date, ID, Total Price, and Status.",
        "Order Confirmation: Success message with Order ID."
    ]
    for layout in layouts:
        document.add_paragraph(layout, style='List Bullet')

    # 11. Source Code
    document.add_heading('11. Source Code', level=1)
    document.add_paragraph("store/models.py (Snippet)").bold = True
    document.add_paragraph(
        "class Order(models.Model):\n"
        "    STATUS_CHOICES = (('Pending', 'Pending'), ...)\n"
        "    user = models.ForeignKey(User, on_delete=models.CASCADE)\n"
        "    total_price = models.DecimalField(max_digits=10, decimal_places=2)\n"
    )

    document.add_paragraph("store/views.py (Checkout Snippet)").bold = True
    document.add_paragraph(
        "def checkout_view(request):\n"
        "    if request.method == 'POST':\n"
        "        order = Order.objects.create(user=request.user, ...)\n"
        "        return redirect('order_confirmation', order_id=order.id)\n"
    )

    # 12. References
    document.add_heading('12. Reference', level=1)
    refs = [
        "Django Documentation: https://docs.djangoproject.com/",
        "Bootstrap Documentation: https://getbootstrap.com/",
        "Python Official Site: https://www.python.org/"
    ]
    for ref in refs:
        document.add_paragraph(ref, style='List Number')

    document.save('CraveCart_Report1.docx')
    print("Document saved successfully as CraveCart_Report.docx")

if __name__ == "__main__":
    main()
